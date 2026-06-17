import os
import tempfile
from datetime import datetime

from django.conf import settings
from django.template.loader import render_to_string
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


from sabe.apps.findings.models import Finding, Evidence
from sabe.apps.legal.models import FindingLegalFramework


def _get_report_data(audit):
    findings = Finding.objects.filter(audit=audit).prefetch_related('evidences__document')
    data = {
        'audit': audit,
        'findings': findings,
        'generated_at': datetime.now().strftime('%d/%m/%Y %H:%M'),
    }
    return data


def generate_docx_report(audit, user):
    data = _get_report_data(audit)
    doc = DocxDocument()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Capa
    for _ in range(6):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TRIBUNAL DE CONTAS DO ESTADO DE RORAIMA')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('SISTEMA DE AUDITORIA BASEADO EM EVIDÊNCIAS')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('')
    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_title.add_run(f'RELATÓRIO DE AUDITORIA')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Auditoria: {audit.numero}/{audit.exercicio}\n')
    info.add_run(f'Unidade Jurisdicionada: {audit.unidade_jurisdicionada}\n')
    info.add_run(f'Responsável: {audit.responsavel}\n')
    info.add_run(f'Data: {data["generated_at"]}')

    doc.add_page_break()

    # Sumário
    doc.add_heading('SUMÁRIO', level=1)
    doc.add_paragraph('1. Introdução')
    doc.add_paragraph('2. Escopo')
    doc.add_paragraph('3. Metodologia')
    doc.add_paragraph('4. Achados de Auditoria')
    doc.add_paragraph('5. Evidências')
    doc.add_paragraph('6. Fundamentação Legal')
    doc.add_paragraph('7. Conclusão')
    doc.add_paragraph('8. Recomendações')

    doc.add_page_break()

    # 1. Introdução
    doc.add_heading('1. INTRODUÇÃO', level=1)
    p = doc.add_paragraph()
    p.add_run(f'O presente relatório apresenta os resultados da auditoria realizada no âmbito do ')
    p.add_run(f'Tribunal de Contas do Estado de Roraima, referente ao processo ')
    p.add_run(f'{audit.processo or "n/a"}.' if audit.processo else 'n/a.')
    doc.add_paragraph(f'Período da auditoria: {audit.data_inicio.strftime("%d/%m/%Y")} a '
                      f'{audit.data_encerramento.strftime("%d/%m/%Y") if audit.data_encerramento else "em andamento"}')
    doc.add_paragraph(f'Objeto: {audit.objeto}')

    # 2. Escopo
    doc.add_heading('2. ESCOPO', level=1)
    doc.add_paragraph(f'A auditoria abrangeu a unidade jurisdicionada {audit.unidade_jurisdicionada}, '
                      f'tendo como responsável {audit.responsavel}.')
    doc.add_paragraph(f'Foram analisados {audit.documents.count()} documentos e identificados '
                      f'{audit.findings.count()} achados de auditoria.')

    # 3. Metodologia
    doc.add_heading('3. METODOLOGIA', level=1)
    doc.add_paragraph('A metodologia adotada compreendeu as seguintes etapas:')
    doc.add_paragraph('a) Análise documental com suporte de OCR e extração de texto;', style='List Bullet')
    doc.add_paragraph('b) Identificação e registro de evidências;', style='List Bullet')
    doc.add_paragraph('c) Classificação de achados segundo critérios de materialidade e relevância;', style='List Bullet')
    doc.add_paragraph('d) Vinculação de fundamentação legal;', style='List Bullet')
    doc.add_paragraph('e) Elaboração de recomendações.', style='List Bullet')

    # 4. Achados
    doc.add_heading('4. ACHADOS DE AUDITORIA', level=1)
    for i, finding in enumerate(data['findings'], 1):
        doc.add_heading(f'4.{i}. {finding.title}', level=2)
        doc.add_paragraph(f'Classificação: {finding.get_classificacao_display()}')
        if finding.description:
            p = doc.add_paragraph()
            p.add_run('Descrição: ').bold = True
            p.add_run(finding.description)
        if finding.criterio:
            p = doc.add_paragraph()
            p.add_run('Critério: ').bold = True
            p.add_run(finding.criterio)
        if finding.condicao:
            p = doc.add_paragraph()
            p.add_run('Condição: ').bold = True
            p.add_run(finding.condicao)
        if finding.causa:
            p = doc.add_paragraph()
            p.add_run('Causa: ').bold = True
            p.add_run(finding.causa)
        if finding.efeito:
            p = doc.add_paragraph()
            p.add_run('Efeito: ').bold = True
            p.add_run(finding.efeito)
        if finding.recomendacao:
            p = doc.add_paragraph()
            p.add_run('Recomendação: ').bold = True
            p.add_run(finding.recomendacao)

        evidences = finding.evidences.select_related('document').all()
        if evidences:
            doc.add_paragraph('')
            p = doc.add_paragraph()
            p.add_run('Evidências:').bold = True
            for ev in evidences:
                doc.add_paragraph(
                    f'- Documento: {ev.document.original_filename}, '
                    f'Página {ev.page_number}',
                    style='List Bullet'
                )

    doc.add_page_break()

    # 5. Evidências
    doc.add_heading('5. EVIDÊNCIAS', level=1)
    all_evidences = Evidence.objects.filter(finding__audit=audit).select_related('document', 'finding')
    for ev in all_evidences:
        doc.add_paragraph(
            f'Documento: {ev.document.original_filename} | '
            f'Página: {ev.page_number} | '
            f'Achado: {ev.finding.title}'
        )
        doc.add_paragraph(f'Texto: {ev.captured_text[:200]}...' if len(ev.captured_text) > 200 else ev.captured_text)
        doc.add_paragraph('')

    # 6. Fundamentação Legal
    doc.add_heading('6. FUNDAMENTAÇÃO LEGAL', level=1)
    found_any = False
    for finding in data['findings']:
        lfs = FindingLegalFramework.objects.filter(
            finding=finding
        ).select_related('legal_framework')
        if lfs:
            found_any = True
            doc.add_heading(f'{finding.title}', level=2)
            for lf in lfs:
                lf_obj = lf.legal_framework
                p = doc.add_paragraph()
                p.add_run(f'{lf_obj.tipo_display()} - {lf_obj.numero}').bold = True
                if lf_obj.ementa:
                    doc.add_paragraph(f'Ementa: {lf_obj.ementa}')
                if lf_obj.texto:
                    doc.add_paragraph(f'Texto: {lf_obj.texto}')
                doc.add_paragraph('')
    if not found_any:
        doc.add_paragraph('Nenhuma fundamentação legal vinculada aos achados.')

    # 7. Conclusão
    doc.add_heading('7. CONCLUSÃO', level=1)
    doc.add_paragraph(f'Foram realizadas as análises pertinentes no âmbito da auditoria '
                      f'{audit.numero}/{audit.exercicio}, tendo sido identificados '
                      f'{audit.findings.count()} achados, suportados por evidências documentais.')
    if audit.observacoes:
        doc.add_paragraph(f'Observações: {audit.observacoes}')

    # 8. Recomendações
    doc.add_heading('8. RECOMENDAÇÕES', level=1)
    recommendations = [f for f in data['findings'] if f.recomendacao]
    if recommendations:
        for i, finding in enumerate(recommendations, 1):
            doc.add_paragraph(f'{i}. {finding.recomendacao}')
    else:
        doc.add_paragraph('Nenhuma recomendação registrada.')

    output_dir = settings.REPORT_OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    filename = f'relatorio_{audit.numero}_{audit.exercicio}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def generate_pdf_report(audit, user):
    from io import BytesIO
    from xhtml2pdf import pisa

    data = _get_report_data(audit)
    html_string = render_to_string('reports/pdf_template.html', data)

    output_dir = settings.REPORT_OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    filename = f'relatorio_{audit.numero}_{audit.exercicio}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'wb') as f:
        pisa_status = pisa.CreatePDF(
            html_string, dest=f,
            encoding='utf-8',
        )

    if pisa_status.err:
        raise RuntimeError(f'Erro ao gerar PDF: {pisa_status.err}')
    return filepath
